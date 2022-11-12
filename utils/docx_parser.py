from docx import Document


class DocxParser:
    """
    Parses .docx files content: paragraphs, images, tables

    Parameters
    ----------
    file_path : str
        Relative path to file

    Attributes
    ----------
    docx : Document
        Document object
    """

    def __init__(self, file_path):
        self.docx = Document(file_path)

    def get_docx_images(self):
        """
        Returns a list of images in the doc.

        Returns
        -------
        images list : list
            List of images in the doc
        """
        images = []
        for image in self.docx.inline_shapes:
            # 3 is wdInlineShapePicture
            # see https://docs.microsoft.com/en-us/office/vba/api/Word.WdInlineShapeType
            if image.type != 3:
                raise ValueError(f'Expected image.type 3, got {image.type} instead.')
            images.append(image.type)

        return images

    def get_docx_text(self):
        """
        Returns a list of docx paragraphs (text).

        Returns
        -------
        paragraphs_list : list
            List of docx paragraphs (text)
        """
        paragraphs_list = []

        for paragraph in self.docx.paragraphs:
            paragraphs_list.append(paragraph.text)

        return paragraphs_list

    def get_docx_table_data(self):
        """
        Returns a list of docx tables data.

        Returns
        -------
        tables_data_list : list
            List of docx tables data
        """
        tables_data_list = []

        for table in self.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        tables_data_list.append(paragraph.text)

        return tables_data_list
